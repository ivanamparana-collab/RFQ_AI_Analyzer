"""
RFQ_AI_Analyzer
================

Este programa lee un archivo PDF de una solicitud de cotizacion (RFQ),
extrae su texto, lo analiza con OpenAI y exporta los resultados a Word y Excel.

Pensado para ejecutarse sin cambios en Google Colab o en un entorno local.
"""

# =============================================================================
# 1. IMPORTACION DE LIBRERIAS
# =============================================================================

import json
import os
from pathlib import Path
from typing import Any, Dict

import pandas as pd
from docx import Document
from openai import OpenAI
from pypdf import PdfReader


# =============================================================================
# 2. CONFIGURACION GENERAL
# =============================================================================

# Nombre del PDF que debe estar en la misma carpeta del script o subirse a Colab.
PDF_ENTRADA = "RFQ_001.pdf"

# Archivos de salida que generara el programa.
WORD_SALIDA = "analisis_rfq.docx"
EXCEL_SALIDA = "analisis_rfq.xlsx"

# Modelo solicitado para el analisis.
MODELO_OPENAI = "gpt-4o-mini"


# =============================================================================
# 3. EXTRACCION DE TEXTO DEL PDF
# =============================================================================

def extraer_texto_pdf(ruta_pdf: str) -> str:
    """
    Extrae el texto de todas las paginas de un archivo PDF usando pypdf.

    Args:
        ruta_pdf: Ruta del archivo PDF de entrada.

    Returns:
        Texto completo extraido del PDF.

    Raises:
        FileNotFoundError: Si el archivo PDF no existe.
        ValueError: Si el PDF no contiene texto legible.
    """
    ruta = Path(ruta_pdf)

    if not ruta.exists():
        raise FileNotFoundError(
            f"No se encontro el archivo {ruta_pdf}. "
            "En Google Colab, sube el PDF antes de ejecutar el script."
        )

    lector = PdfReader(str(ruta))
    textos_paginas = []

    for numero_pagina, pagina in enumerate(lector.pages, start=1):
        texto_pagina = pagina.extract_text() or ""
        texto_pagina = texto_pagina.strip()

        if texto_pagina:
            textos_paginas.append(f"--- Pagina {numero_pagina} ---\n{texto_pagina}")

    texto_completo = "\n\n".join(textos_paginas).strip()

    if not texto_completo:
        raise ValueError(
            "No se pudo extraer texto del PDF. "
            "Si el RFQ esta escaneado como imagen, primero debe aplicarse OCR."
        )

    return texto_completo


# =============================================================================
# 4. ANALISIS DEL RFQ CON OPENAI
# =============================================================================

def crear_prompt_analisis(texto_rfq: str) -> str:
    """
    Construye el prompt que se enviara al modelo de OpenAI.

    Args:
        texto_rfq: Texto extraido del PDF.

    Returns:
        Prompt completo con instrucciones y contenido del RFQ.
    """
    return f"""
Analiza el siguiente documento RFQ y devuelve exclusivamente un JSON valido.

El JSON debe tener exactamente estas claves:
- alcance_del_proyecto
- entregables
- fechas_importantes
- riesgos
- informacion_faltante
- supuestos
- resumen_ejecutivo

Reglas:
- Escribe en espanol.
- Usa listas de strings para todas las claves excepto resumen_ejecutivo.
- resumen_ejecutivo debe ser un string breve y claro.
- Si una seccion no esta explicita en el RFQ, indicalo como informacion no especificada.
- No inventes datos; separa claramente los supuestos.

Texto del RFQ:
{texto_rfq}
"""


def analizar_rfq_con_openai(texto_rfq: str) -> Dict[str, Any]:
    """
    Envia el texto del RFQ a GPT-4o-mini y obtiene un analisis estructurado.

    La API key debe estar disponible en la variable de entorno OPENAI_API_KEY.
    En Google Colab se puede configurar con:
        import os
        os.environ["OPENAI_API_KEY"] = "tu_api_key"

    Args:
        texto_rfq: Texto completo extraido del PDF.

    Returns:
        Diccionario con el analisis del RFQ.

    Raises:
        EnvironmentError: Si no se encuentra la variable OPENAI_API_KEY.
        ValueError: Si la respuesta del modelo no es un JSON valido.
    """
    if not os.getenv("OPENAI_API_KEY"):
        raise EnvironmentError(
            "Falta la variable de entorno OPENAI_API_KEY. "
            "Configura tu API key antes de ejecutar el analisis."
        )

    cliente = OpenAI()
    prompt = crear_prompt_analisis(texto_rfq)

    respuesta = cliente.chat.completions.create(
        model=MODELO_OPENAI,
        messages=[
            {
                "role": "system",
                "content": (
                    "Eres un analista experto en RFQ, licitaciones, "
                    "alcances de proyecto, riesgos y entregables."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        response_format={"type": "json_object"},
        temperature=0.2,
    )

    contenido = respuesta.choices[0].message.content

    try:
        return json.loads(contenido)
    except json.JSONDecodeError as error:
        raise ValueError(
            "La respuesta del modelo no pudo interpretarse como JSON valido."
        ) from error


# =============================================================================
# 5. EXPORTACION A WORD
# =============================================================================

def exportar_a_word(analisis: Dict[str, Any], ruta_salida: str) -> None:
    """
    Exporta el analisis del RFQ a un archivo Word usando python-docx.

    Args:
        analisis: Diccionario con el analisis generado por OpenAI.
        ruta_salida: Ruta del archivo .docx que se creara.
    """
    documento = Document()

    documento.add_heading("Analisis de RFQ", level=1)
    documento.add_paragraph(analisis.get("resumen_ejecutivo", "Sin resumen ejecutivo."))

    secciones = [
        ("Alcance del proyecto", "alcance_del_proyecto"),
        ("Entregables", "entregables"),
        ("Fechas importantes", "fechas_importantes"),
        ("Riesgos", "riesgos"),
        ("Informacion faltante", "informacion_faltante"),
        ("Supuestos", "supuestos"),
    ]

    for titulo, clave in secciones:
        documento.add_heading(titulo, level=2)
        valores = analisis.get(clave, [])

        if isinstance(valores, str):
            valores = [valores]

        if not valores:
            valores = ["Informacion no especificada."]

        for valor in valores:
            documento.add_paragraph(str(valor), style="List Bullet")

    documento.save(ruta_salida)


# =============================================================================
# 6. EXPORTACION A EXCEL
# =============================================================================

def exportar_a_excel(analisis: Dict[str, Any], ruta_salida: str) -> None:
    """
    Exporta el analisis del RFQ a un archivo Excel usando pandas.

    Args:
        analisis: Diccionario con el analisis generado por OpenAI.
        ruta_salida: Ruta del archivo .xlsx que se creara.
    """
    filas = []

    for seccion, contenido in analisis.items():
        if isinstance(contenido, list):
            for elemento in contenido:
                filas.append({"seccion": seccion, "contenido": elemento})
        else:
            filas.append({"seccion": seccion, "contenido": contenido})

    dataframe = pd.DataFrame(filas)
    dataframe.to_excel(ruta_salida, index=False)


# =============================================================================
# 7. FLUJO PRINCIPAL DEL PROGRAMA
# =============================================================================

def main() -> None:
    """
    Ejecuta el flujo completo:
    1. Leer PDF.
    2. Extraer texto.
    3. Analizar con OpenAI.
    4. Exportar a Word.
    5. Exportar a Excel.
    """
    print("Iniciando analisis del RFQ...")

    texto_rfq = extraer_texto_pdf(PDF_ENTRADA)
    print("Texto extraido correctamente del PDF.")

    analisis = analizar_rfq_con_openai(texto_rfq)
    print("Analisis generado correctamente con OpenAI.")

    exportar_a_word(analisis, WORD_SALIDA)
    print(f"Archivo Word creado: {WORD_SALIDA}")

    exportar_a_excel(analisis, EXCEL_SALIDA)
    print(f"Archivo Excel creado: {EXCEL_SALIDA}")

    print("Proceso finalizado.")


if __name__ == "__main__":
    main()
